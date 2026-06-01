#!/usr/bin/env python3
"""Generate a Word (.docx) summary of the jasper-deploy skill's capabilities."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x34, 0x49, 0x5E)
GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# base font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


# ---- Title ----------------------------------------------------------------
title = doc.add_heading("jasper-deploy", level=0)
for r in title.runs:
    r.font.color.rgb = NAVY
sub = doc.add_paragraph("Capabilities Summary — a Claude Code skill for designing, "
                        "compiling, and deploying JasperReports artifacts to "
                        "JasperReports Server")
sub.runs[0].italic = True
sub.runs[0].font.color.rgb = GREY
meta = doc.add_paragraph("Generated 2026-06-01  ·  repo: AI-JasperReports-Generator  "
                         "·  location: .claude/skills/jasper-deploy/")
meta.runs[0].font.size = Pt(8.5)
meta.runs[0].font.color.rgb = GREY

# ---- Overview -------------------------------------------------------------
heading("Overview", 1)
doc.add_paragraph(
    "jasper-deploy automates the full JasperReports 7 pipeline against a "
    "PostgreSQL/PostGIS database and a JasperReports Server (Pro) instance: "
    "scaffold a report from a SQL query, compile and validate it locally, "
    "create the data source, deploy it over the REST v2 API, and run it to PDF. "
    "It covers tabular reports plus a full range of visualizations — from "
    "open-source JFreeChart charts to commercial HTML5 (HighCharts) charts and "
    "FusionMaps choropleths.")

# ---- Pipeline -------------------------------------------------------------
heading("The pipeline", 1)
table(
    ["Stage", "Script", "What it does"],
    [
        ["1. Design", "scaffold_jrxml.py",
         "Introspects a SQL query's columns (via psql) and emits a JR7-native "
         ".jrxml: title, column header, detail band, page footer. Optional "
         "--chart adds a JFreeChart in the summary band."],
        ["2. Compile", "compile_jrxml.ps1 + CompileReport.java",
         "Compiles .jrxml → .jasper with JDK 11 single-file source launch "
         "against the JR 7.0.6 runtime; also the fastest JR7-validity check."],
        ["3. Data source", "create_datasource.ps1",
         "Creates/updates a JDBC data source on the server via REST v2 "
         "(PostgreSQL driver ships with JRS)."],
        ["4. Deploy", "deploy_report.ps1",
         "Wraps the .jrxml as a reportUnit (base64-inlined) and PUTs it to "
         "/rest_v2/resources. -Overwrite handles re-deploys."],
        ["5. Run / preview", "(REST + RenderPng.java / pypdfium2)",
         "Runs the report to PDF server-side; renders a page to PNG locally for "
         "visual checks."],
    ],
    widths=[1.0, 2.0, 3.5])

# ---- Visualization: community ---------------------------------------------
heading("Visualization — community (compile + preview locally, then deploy)", 1)
doc.add_paragraph(
    "Rendered by the open-source JasperReports library; the runtime classpath "
    "includes the charts, barcode4j, JFreeChart, and ZXing jars.")
table(
    ["Capability", "jrxml form", "Example report"],
    [
        ["JFreeChart charts — pie, bar, 3D bar, line, area, stacked",
         'kind="chart" chartType="…" + dataset (pie/category) + plot',
         "metro_population_piechart, metro_population_bar"],
        ["Spider / radar chart",
         'component kind="spiderChart" (series / category / value)',
         "metro_population_spider"],
        ["Barcodes & QR — QR Code, DataMatrix, Code128, …",
         'component kind="barcode4j:QRCode|DataMatrix|Code128"',
         "barcode_demo"],
        ["Tabular reports (single + grouped, page X of Y)",
         "title / columnHeader / detail / group bands",
         "county_summary, tx_addr_zip_summary, tx_featnames_top_streets"],
    ],
    widths=[2.2, 2.8, 1.5])

# ---- Visualization: pro ---------------------------------------------------
heading("Visualization — Pro (server-rendered; deploy → run to validate)", 1)
doc.add_paragraph(
    "Commercial JasperReports Server components, authored in the legacy 6.x "
    "jrxml format and rendered server-side (the open-source library cannot "
    "compile them locally).")
table(
    ["Capability", "jrxml form", "Example report"],
    [
        ["HTML5 charts (HighCharts) — column, stacked, donut, spline, "
         "heatmap, treemap, solid gauge",
         'hc:chart type="…" + multiAxisData (row buckets + measures)',
         "metro_population_html5"],
        ["FusionMaps choropleth — colour regions by value",
         "fm:map + mapNameExpression + colorRange(s) + entity (id + value)",
         "tx_county_density_map (Texas counties)"],
        ["Also available on the server",
         "Fusion charts / gauges / widgets; Ad Hoc views & dashboards (web UI)",
         "—"],
    ],
    widths=[2.2, 2.8, 1.5])
note = doc.add_paragraph()
note.add_run("Note: ").bold = True
note.add_run("the installed Texas FusionMap is keyed by county FIPS, so a county "
             "choropleth binds idExpression directly to (countyfp::int)::text — "
             "no lookup table needed.")
note.runs[-1].font.color.rgb = GREY

# ---- Environment ----------------------------------------------------------
heading("Environment & toolchain", 1)
bullets([
    "JasperReports Server PRO/Enterprise 10.0.0 at http://localhost:8081/"
    "jasperserver-pro (REST v2, HTTP Basic). Note: port 8080 is an unrelated "
    "Bearer-gated service — not JRS.",
    "JasperReports 7.0.6 runtime in C:\\Users\\rgorsuch\\jasperreports-lib "
    "(+ charts, barcode4j, JFreeChart, ZXing jars).",
    "JDK 11 (single-file source launch), psql 14, curl 8.x, Python 3.14.",
    "Database: PostgreSQL/PostGIS postgis_34_sample (TIGER geocoder data).",
    "Credentials load from script params → env vars (JRS_URL/USER/PASS) → a "
    "gitignored jrs.config.json.",
])

# ---- Operational notes ----------------------------------------------------
heading("Key operational notes", 1)
bullets([
    "JR7 jrxml is NOT 6.x-compatible (no namespace; <query>; elements use "
    'kind="…" with flattened geometry). See references/jr7-schema.md.',
    "JRS SQL validator rejects queries that don't begin with SELECT — rewrite "
    "CTEs (WITH) as nested subqueries; window functions are fine.",
    'A chart/map bound to the main dataset in the title band needs '
    'evaluationTime="Report", or it renders blank.',
    "Re-deploying an existing report needs -Overwrite (optimistic-locking 409).",
    "Preview Pro (server-only) reports by rasterizing the server PDF with "
    "pypdfium2.",
])

out = r"C:\Users\rgorsuch\tx-geocoder\jasper-deploy-capabilities.docx"
doc.save(out)
print("saved", out)
